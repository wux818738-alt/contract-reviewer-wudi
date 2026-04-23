#!/usr/bin/env python3
"""
pdf_ocr.py — 合同审查 SKILL 的 PDF 扫描件处理模块

功能：
  1. 自动检测 PDF 是文本型还是扫描件
  2. 扫描件 → Tesseract OCR → 可搜索 PDF / 文本提取 / docx
  3. 输出段落结构化文本，供 contract_parser.py / apply_changes.py 使用

依赖：
  - tesseract (brew install tesseract tesseract-lang)
  - pytesseract (pip3 install pytesseract)
  - pdf2image (pip3 install pdf2image)
  - PyMuPDF (pip3 install pymupdf)
  - Pillow (pip3 install pillow)
  - numpy
"""
from __future__ import annotations

import re, json, sys, os, tempfile, shutil
from pathlib import Path
from typing import Optional
from dataclasses import dataclass, field

# ── 依赖检测 ────────────────────────────────────────────────────────────────

def _check_deps():
    """检查运行时依赖，返回缺失列表"""
    missing = []
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
    except Exception:
        missing.append('pytesseract')
    try:
        import fitz
    except Exception:
        missing.append('pymupdf')
    try:
        import pdf2image
    except Exception:
        missing.append('pdf2image')
    try:
        from PIL import Image
    except Exception:
        missing.append('pillow')
    try:
        import numpy as np
    except Exception:
        missing.append('numpy')
    return missing


def install_deps():
    """自动安装缺失依赖"""
    import subprocess, sys
    deps = _check_deps()
    if not deps:
        return True
    print(f"[安装] 缺少依赖: {', '.join(deps)}")
    try:
        subprocess.run([sys.executable, '-m', 'pip', 'install', '-q'] + deps,
                       check=True)
        print("[安装] 依赖安装完成")
        return True
    except Exception as e:
        print(f"[安装] 失败: {e}")
        return False


# ── 数据结构 ────────────────────────────────────────────────────────────────

@dataclass
class OcrResult:
    """OCR 结果"""
    text: str                          # 完整 OCR 文本
    paragraphs: list[dict] = field(default_factory=list)
    # [{'index': int, 'text': str, 'bbox': (x0,y0,x1,y1), 'confidence': float}]
    is_scanned: bool = True            # 是否为扫描件
    page_count: int = 0
    source_file: str = ''


# ── 核心函数 ───────────────────────────────────────────────────────────────

def detect_if_scanned(pdf_path: str) -> bool:
    """判断 PDF 是文本型还是扫描图像"""
    import fitz
    doc = fitz.open(pdf_path)
    try:
        for page in doc:
            text = page.get_text('text', flags=0).strip()
            if len(text) > 50:
                return False  # 有可提取文字 → 文本型 PDF
            # 检查是否有嵌入图片（扫描特征）
            img_list = page.get_images(full=True)
            for img in img_list:
                if img[2] > 10000:  # 图片面积 > 10000 sq pt
                    return True     # 大图 = 扫描特征
    finally:
        doc.close()
    return True  # 没文字也没大图 → 默认当扫描件处理


def ocr_pdf_pages(pdf_path: str,
                  lang: str = 'chi_sim+eng',
                  dpi: int = 300,
                  ) -> tuple[list[dict], str]:
    """
    对 PDF 的每一页做 OCR，返回段落列表 + 完整文本
    lang: tesseract 语言代码
      chi_sim  = 简体中文
      eng      = 英文
      chi_sim+eng = 混排
    """
    import fitz
    import pytesseract
    from PIL import Image
    import numpy as np

    doc = fitz.open(pdf_path)
    paragraphs: list[dict] = []
    full_text_parts: list[str] = []
    page_count = len(doc)

    for page_num in range(page_count):
        page = doc[page_num]

        # ── 方案 A：直接用 PyMuPDF 渲染 Pixmap ───────────────
        try:
            # 300 DPI 渲染
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)

            # Pixmap → PIL Image
            img_data = pix.samples
            img = Image.frombytes('RGB', [pix.width, pix.height], img_data)

            # ── Tesseract OCR ──────────────────────────────────
            # 获取带位置信息的文本块
            data = pytesseract.image_to_data(
                img, lang=lang, output_type=pytesseract.Output.DICT,
                config='--psm 4'   # 假设单栏文本
            )

            # 解析数据块，过滤低置信度（<60）
            n_boxes = len(data['text'])
            page_paras = []
            current_para = []
            current_conf = []

            for i in range(n_boxes):
                word = data['text'][i].strip()
                conf = int(data['conf'][i])

                if conf < 60 or not word:
                    # 置信度低或空，跳过但可能结束当前段
                    if current_para:
                        page_paras.append(' '.join(current_para))
                        current_para = []
                    continue

                current_para.append(word)
                current_conf.append(conf)

            if current_para:
                page_paras.append(' '.join(current_para))

            # 计算全局段落索引偏移
            offset = len(paragraphs)
            for idx, para_text in enumerate(page_paras):
                if len(para_text) < 3:
                    continue
                paragraphs.append({
                    'page': page_num + 1,
                    'index': offset + len([p for p in paragraphs if p.get('page') == page_num + 1]),
                    'text': para_text,
                    'bbox': None,  # 可扩展：加入 bbox
                    'confidence': round(sum(current_conf) / len(current_conf), 1)
                        if current_conf else 0,
                })
                full_text_parts.append(para_text)

        except Exception as e:
            print(f"[OCR] 第 {page_num+1} 页失败: {e}")
            continue

    doc.close()
    return paragraphs, '\n'.join(full_text_parts)


def convert_to_searchable_pdf(src_pdf: str,
                              out_pdf: str,
                              lang: str = 'chi_sim+eng',
                              dpi: int = 300) -> bool:
    """
    将扫描 PDF 转换为可搜索 PDF（文本叠加在图片上）
    返回是否成功
    """
    import fitz
    import pytesseract
    from PIL import Image
    import numpy as np

    try:
        doc = fitz.open(src_pdf)
        out_doc = fitz.open()

        for page_num in range(len(doc)):
            page = doc[page_num]

            # 渲染为高分辨率图片
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
            img_data = pix.samples
            img = Image.frombytes('RGB', [pix.width, pix.height], img_data)

            # OCR 获取文本 + 位置
            data = pytesseract.image_to_data(
                img, lang=lang, output_type=pytesseract.Output.DICT,
                config='--psm 4')

            # 创建新页面
            width = pix.width / dpi * 72   # 还原为点（72 DPI）
            height = pix.height / dpi * 72
            out_page = out_doc.new_page(width=width, height=height)

            # 插入图片作为背景
            img_path = Path(tempfile.gettempdir()) / f'ocr_page_{page_num}.png'
            img.save(img_path)
            out_page.insert_image(img_path, keep_proportion=False)

            # 在图片上叠加不可见的文字层（便于 PDF 搜索引擎索引）
            n_boxes = len(data['text'])
            for i in range(n_boxes):
                word = data['text'][i].strip()
                conf = int(data['conf'][i])
                if conf < 60 or not word:
                    continue
                x = data['left'][i] / dpi * 72
                y = data['top'][i] / dpi * 72
                w = data['width'][i] / dpi * 72
                h = data['height'][i] / dpi * 72
                out_page.insert_text(
                    (x, y + h * 0.8),
                    word,
                    fontsize=max(6, h * 0.7),
                    color=(0, 0, 0),
                    render_mode=3,  # 不可见文字层
                )

            Path(img_path).unlink(missing_ok=True)

        out_doc.save(out_pdf)
        out_doc.close()
        doc.close()
        return True

    except Exception as e:
        print(f"[可搜索PDF] 失败: {e}")
        return False


def convert_scan_to_docx(src_pdf: str,
                          out_docx: str,
                          lang: str = 'chi_sim+eng',
                          dpi: int = 300) -> bool:
    """
    将扫描 PDF 转换为 .docx（文本内容），供后续 SKILL 流程处理
    """
    import fitz, pytesseract
    from PIL import Image
    import zipfile, io
    import xml.etree.ElementTree as ET
    from pathlib import Path

    doc = fitz.open(src_pdf)
    all_text: list[str] = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        mat = fitz.Matrix(dpi / 72, dpi / 72)
        pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
        img = Image.frombytes('RGB', [pix.width, pix.height], pix.samples)

        # 整页 OCR，保留段落结构
        text = pytesseract.image_to_string(img, lang=lang, config='--psm 4')
        # 清理常见 OCR 错误
        text = _clean_ocr_text(text)
        all_text.append(text)

    doc.close()

    # 生成最小化 docx（纯文本）
    _make_plain_docx(out_docx, all_text)
    return True


# ── 辅助函数 ───────────────────────────────────────────────────────────────

def _clean_ocr_text(text: str) -> str:
    """清理 OCR 常见错误"""
    # 常见数字/字符混淆
    replacements = [
        (r'[（(]\d+[)）](?=[一二三四五六七八九十])', ''),  # 去掉 "(1)一" 的括号
        (r'[—―–-]{2,}', '—'),                               # 合并多个破折号
        (r'[ |　]{2,}', ' '),                               # 多个空格
        (r'([\u4e00-\u9fff])[|Il1]([\u4e00-\u9fff])', r'\1I\2'),  # 竖线误识为 I/l
    ]
    for pat, repl in replacements:
        text = re.sub(pat, repl, text)
    # 去掉每行首尾空白
    lines = [ln.strip() for ln in text.splitlines()]
    return '\n'.join(ln for ln in lines if ln)


def _make_plain_docx(out_path: str, paragraphs: list[str]):
    """用 python-docx 生成纯文本 docx（如果安装了的话）"""
    try:
        from docx import Document
        doc = Document()
        for para in paragraphs:
            if not para.strip():
                continue
            for line in para.split('\n'):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)
        doc.save(out_path)
        return
    except ImportError:
        pass

    # 兜底：用 PyMuPDF 生成 docx（XML 直接写入）
    _make_docx_xml(out_path, paragraphs)


def _make_docx_xml(out_path: str, paragraphs: list[str]):
    """不依赖 python-docx，用 zipfile+XML 生成最简 docx"""
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'

    def qn(tag):
        return f'{{{W}}}{tag}'

    body_para = ''
    for text in paragraphs:
        safe = (text.replace('&', '&amp;')
                      .replace('<', '&lt;')
                      .replace('>', '&gt;'))
        body_para += (
            f'<w:p>'
            f'<w:r><w:t xml:space="preserve">{safe}</w:t></w:r>'
            f'</w:p>'
        )

    doc_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document {NS}>
<w:body>
{body_para}
<w:sectPr>
  <w:pgSz w:w="12240" w:h="15840"/>
  <w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/>
</w:sectPr>
</w:body>
</w:document>'''

    rels = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
</Relationships>'''

    ct = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>'''

    with zipfile.ZipFile(out_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('[Content_Types].xml', ct)
        zf.writestr('_rels/.rels', rels)
        zf.writestr('word/_rels/document.xml.rels', rels)
        zf.writestr('word/document.xml', doc_xml)


# ── 主入口 ─────────────────────────────────────────────────────────────────

def process_pdf(pdf_path: str,
                output_dir: Optional[str] = None,
                lang: str = 'chi_sim+eng',
                dpi: int = 300) -> OcrResult:
    """
    主入口：对 PDF 执行完整 OCR 处理流程

    参数:
        pdf_path: PDF 文件路径
        output_dir: 输出目录（默认与 PDF 同目录）
        lang: Tesseract 语言代码，默认 chi_sim+eng
        dpi: 渲染分辨率，默认 300

    返回:
        OcrResult 对象（含 paragraphs、text、is_scanned）

    输出文件（写入 output_dir）:
        {原名}_ocr.txt         — 纯 OCR 文本
        {原名}_ocr_docx.docx   — 可编辑 docx（供后续 SKILL 处理）
        {原名}_searchable.pdf  — 可搜索 PDF（可选）
    """
    missing = _check_deps()
    if missing:
        if not install_deps():
            raise RuntimeError(f"缺少依赖: {', '.join(missing)}，请手动安装后重试")

    pdf_path = str(Path(pdf_path).resolve())
    if not Path(pdf_path).exists():
        raise FileNotFoundError(f"文件不存在: {pdf_path}")

    out_dir = Path(output_dir) if output_dir else Path(pdf_path).parent
    stem = Path(pdf_path).stem

    # ── 1. 检测类型 ───────────────────────────────────────
    is_scanned = detect_if_scanned(pdf_path)
    print(f"[PDF OCR] {'扫描件' if is_scanned else '文本型PDF'}: {Path(pdf_path).name}")

    # ── 2. 文本型 PDF：直接提取文字 ────────────────────────
    if not is_scanned:
        import fitz
        doc = fitz.open(pdf_path)
        all_text_parts = []
        paragraphs = []
        for page_num, page in enumerate(doc):
            page_text = page.get_text('text').strip()
            for line in page_text.splitlines():
                line = line.strip()
                if line:
                    all_text_parts.append(line)
                    paragraphs.append({
                        'page': page_num + 1,
                        'index': len(paragraphs),
                        'text': line,
                        'bbox': None,
                        'confidence': 100.0,
                    })
        doc.close()
        full_text = '\n'.join(all_text_parts)
        result = OcrResult(text=full_text, paragraphs=paragraphs,
                           is_scanned=False, page_count=len(doc),
                           source_file=pdf_path)
    else:
        # ── 3. 扫描件：Tesseract OCR ───────────────────────
        paragraphs, full_text = ocr_pdf_pages(pdf_path, lang=lang, dpi=dpi)
        page_count = len(fitz.open(pdf_path))
        result = OcrResult(text=full_text, paragraphs=paragraphs,
                           is_scanned=True, page_count=page_count,
                           source_file=pdf_path)

    # ── 4. 输出文件 ────────────────────────────────────────
    # OCR 文本
    txt_path = out_dir / f'{stem}_ocr.txt'
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(full_text)
    print(f"  文本保存: {txt_path}")

    # 可编辑 docx（供 contract_parser.py 使用）
    docx_path = out_dir / f'{stem}_ocr.docx'
    try:
        convert_scan_to_docx(pdf_path, str(docx_path), lang=lang, dpi=dpi)
        print(f"  DOCX 保存: {docx_path}")
    except Exception as e:
        print(f"  DOCX 保存失败: {e}")

    return result


# ── CLI ────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(
        description='PDF 扫描件 OCR 处理（支持文本型 PDF 直接提取）')
    parser.add_argument('pdf', help='输入 PDF 文件路径')
    parser.add_argument('-o', '--output-dir', default=None,
                        help='输出目录（默认与 PDF 同目录）')
    parser.add_argument('-l', '--lang', default='chi_sim+eng',
                        help='Tesseract 语言代码，默认 chi_sim+eng')
    parser.add_argument('--dpi', type=int, default=300,
                        help='渲染 DPI，默认 300')
    args = parser.parse_args()

    result = process_pdf(args.pdf, args.output_dir, args.lang, args.dpi)
    print(f"\n✅ 完成：{len(result.paragraphs)} 段落，{len(result.text)} 字")
    print(f"   扫描件: {'是' if result.is_scanned else '否'}")
